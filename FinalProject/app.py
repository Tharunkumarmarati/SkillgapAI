import streamlit as st
import pandas as pd
import numpy as np
import re
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import plotly.graph_objects as go
import plotly.express as px
from io import BytesIO
from reportlab.lib.styles import getSampleStyleSheet
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.platypus import Image as RLImage
import plotly.io as pio
from sentence_transformers import SentenceTransformer
import streamlit.components.v1 as components



# --------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------
st.set_page_config(
    page_title="SkillGap AI",
    page_icon="📊",
    layout="wide"
)


st.title("📊 Skill Gap Analysis Interface")
st.caption("End-to-End Resume vs Job Description Skill Intelligence System")

# --------------------------------------------------
# SKILL DATABASE
# --------------------------------------------------


TECH_SKILLS = [
    "python","java","c","c++","c#","javascript","typescript","go","rust","php","r",
    "html","css","react","angular","vue","node","express","nextjs","spring","spring boot",
    "django","flask","fastapi","rest api","graphql","jwt authentication",
    "sql","mysql","postgresql","oracle","sqlite","mongodb","redis","cassandra","firebase",
    "data analysis","data analytics","data visualization","eda",
    "data cleaning","data preprocessing","data transformation",
    "statistics","descriptive statistics","hypothesis testing",
    "mean","median","variance","confidence interval","python","java","c","c++","c#","javascript","typescript",
    "go","rust","php","r","matlab","scala","kotlin","swift",
    "objective c","dart","groovy","perl","lua","haskell""html","css","sass","less","bootstrap","tailwind css",
    "react","angular","vue","svelte","nextjs","nuxtjs",
    "node.js","express","nest.js",
    "spring","spring boot","django","flask","fastapi",
    "rest api","graphql","grpc","web sockets""sql","mysql","postgresql","oracle","sqlite","mariadb",
    "mongodb","redis","cassandra","couchdb","dynamodb",
    "firebase","elasticsearch","neo4j","snowflake","bigquery","data analysis","data analytics","data science",
    "statistics","probability","hypothesis testing",
    "eda","data cleaning","data preprocessing","data transformation",
    "numpy","pandas","scipy",
    "matplotlib","seaborn","plotly",
    "excel","advanced excel","power bi","tableau","looker","qlik",
    "dashboards","kpis","reporting","business intelligence","machine learning","deep learning","artificial intelligence",
    "supervised learning","unsupervised learning",
    "nlp","computer vision","time series","recommendation systems",
    "reinforcement learning",
    "scikit-learn","tensorflow","keras","pytorch",
    "xgboost","lightgbm","catboost",
    "transformers","bert","gpt","llm","prompt engineering","aws","azure","gcp","cloud computing",
    "docker","kubernetes","helm",
    "ci/cd","jenkins","github actions","gitlab ci","bitbucket pipelines",
    "terraform","ansible","cloudformation",
    "linux","unix","bash","shell scripting","hadoop","spark","pyspark","flink",
    "kafka","airflow","luigi",
    "etl","elt","data pipelines","data ingestion",
    "data warehousing","lakehouse","delta lake","object oriented programming","functional programming",
    "data structures","algorithms",
    "system design","low level design","high level design",
    "microservices","monolithic architecture",
    "design patterns","clean architecture",
    "unit testing","integration testing","system testing",
    "api testing","automation testing","selenium","cypress","application security","web security","network security",
    "encryption","ssl","tls","oauth","jwt",
    "penetration testing","vulnerability assessment","git","github","gitlab","bitbucket",
    "jira","confluence","trello","asana",
    "postman","swagger","soap ui",
    "vscode","intellij","eclipse","pycharm",
    "jupyter","google colab",
    "iam","zero trust","firewalls",

    "sql","mysql","joins","inner join","left join","right join",
    "subqueries","aggregations","group by","having","window functions",

    "numpy","pandas","scipy",
    "matplotlib","seaborn","plotly",
    "power bi","tableau","excel",
    "pivot tables","vlookup","power query",
    "dax","dashboards","kpis","reporting"
    "machine learning","deep learning","nlp","computer vision","reinforcement learning",
    "tensorflow","keras","pytorch","scikit-learn","xgboost","lightgbm",
    "transformers","bert","llm","prompt engineering",
    "aws","azure","gcp","docker","kubernetes","jenkins","git","github actions",
    "terraform","ansible","ci/cd","linux","shell scripting",
    "hadoop","spark","pyspark","kafka","airflow","data pipelines","etl",
    "object oriented programming","data structures","algorithms","system design",
    "microservices","design patterns","software testing","unit testing",
    "integration testing","debugging",
    "oauth","ssl","encryption","application security","web security",
    "postman","swagger","jira","confluence","vscode","intellij","eclipse",
    "github","gitlab","bitbucket","api",
    "jupyter notebook","google colab","visual studio code",
    "git","bitbucket"

]



SOFT_SKILLS = [
    "communication","verbal communication","written communication","communication","verbal communication","written communication",
    "presentation skills","public speaking","stakeholder communication",
    "teamwork","collaboration","cross functional collaboration",
    "leadership","people management","mentoring","coaching",
    "problem solving","analytical thinking","critical thinking",
    "decision making","strategic thinking",
    "time management","prioritization","multitasking",
    "adaptability","flexibility","learning mindset",
    "creativity","innovation","design thinking",
    "conflict resolution","negotiation",
    "emotional intelligence","stress management",
    "attention to detail","ownership","accountability",
    "self motivation","work ethic","initiative",
    "customer focus","business acumen","ethics",
    "leadership","teamwork","collaboration","problem solving",
    "critical thinking","analytical thinking","decision making",
    "time management","prioritization","adaptability","flexibility",
    "creativity","innovation","conflict resolution","negotiation",
    "presentation skills","public speaking","active listening",
    "emotional intelligence","stress management","work ethic",
    "accountability","ownership","attention to detail",
    "learning mindset","self motivation","mentorship","analytical thinking","problem solving","communication",
    "verbal communication","written communication",
    "team collaboration","teamwork",
    "decision making","critical thinking",
    "presentation skills","stakeholder communication",
    "attention to detail","time management"
]

UPSKILL_MAP = {
    "Aws": {
        "title": "AWS Cloud Services",
        "description": "Complete AWS Certified Solutions Architect course",
        "priority": "High"
    },
    "Sql": {
        "title": "SQL for Data Analytics",
        "description": "Learn advanced SQL queries, joins, and optimization",
        "priority": "High"
    },
    "Tensorflow": {
        "title": "Deep Learning with TensorFlow",
        "description": "Hands-on TensorFlow & neural networks training",
        "priority": "Medium"
    },
    "Data Analysis": {
        "title": "Advanced Data Analysis",
        "description": "Statistics, EDA, and real-world data projects",
        "priority": "Medium"
    },
    "Leadership": {
        "title": "Leadership & Management Skills",
        "description": "Develop leadership, decision-making, and people skills",
        "priority": "Medium"
    }
}


# --------------------------------------------------
# HELPERS
# --------------------------------------------------
@st.cache_resource
def load_bert_model():
    return SentenceTransformer("all-MiniLM-L6-v2")

bert_model = load_bert_model()


def extract_text(file):
    if file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        return " ".join(p.extract_text() for p in reader.pages if p.extract_text())
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return " ".join(p.text for p in doc.paragraphs if p.text)
    else:
        return file.read().decode("utf-8")

def clean_text(text):
    text = text.lower()
    text = re.sub(r"[^\w\s]", " ", text)
    return re.sub(r"\s+", " ", text)

def extract_skills(text):
    found = []

    # Longer skills first (CRITICAL)
    all_skills = sorted(
        TECH_SKILLS + SOFT_SKILLS,
        key=len,
        reverse=True
    )

    pattern = r"\b(" + "|".join(map(re.escape, all_skills)) + r")\b"

    for match in re.finditer(pattern, text):
        skill = match.group(1).title()
        if skill not in found:
            found.append(skill)

    return found


def classify(skills):
    tech, soft = [], []
    for s in skills:
        if s.lower() in TECH_SKILLS:
            tech.append(s)
        else:
            soft.append(s)
    return tech, soft

def similarity(resume, jd):
    tfidf = TfidfVectorizer()
    mat = tfidf.fit_transform([resume, jd])
    return round(cosine_similarity(mat[0:1], mat[1:2])[0][0] * 100)

def skill_badges(skills, color):
    if not skills:
        return "<i>No skills found</i>"

    html = "<div style='display:flex; flex-wrap:wrap; gap:6px; overflow-x:auto;'>"

    for s in skills:
        html += f"""
        <span style="
            background-color:{color};
            color:white;
            padding:6px 12px;
            border-radius:20px;
            display:inline-block;
            font-size:14px;
            white-space:nowrap;
        ">
            {s}
        </span>
        """

    html += "</div>"
    return html



def comma_list(skills):
    return ", ".join(sorted(skills)) if skills else "None"

def bert_skill_match(resume_skills, jd_skills,
                   exact_th=0.75, partial_th=0.4):

    results = []

    if not resume_skills or not jd_skills:
        return results

    resume_emb = bert_model.encode(resume_skills, normalize_embeddings=True)
    jd_emb = bert_model.encode(jd_skills, normalize_embeddings=True)

    for i, jd_skill in enumerate(jd_skills):
        scores = cosine_similarity(
            [jd_emb[i]],
            resume_emb
        )[0]

        best_score = float(np.max(scores))

        if best_score >= exact_th:
            status = "Matched"
        elif best_score >= partial_th:
            status = "Partial"
        else:
            status = "Missing"

        results.append({
            "Skill": jd_skill,
            "Similarity": round(best_score * 100, 2),
            "Status": status
        })

    return results

def generate_upskilling_recommendations(missing, partial):
    recommendations = []

    for skill in missing:
        key = skill.title()
        if key in UPSKILL_MAP:
            rec = UPSKILL_MAP[key].copy()
            rec["skill"] = skill
            rec["status"] = "Missing"
            rec["priority"] = "High"
            recommendations.append(rec)
        else:
            recommendations.append({
                "skill": skill,
                "title": f"Learn {skill}",
                "description": f"Build strong foundation in {skill} to meet job requirements",
                "priority": "High",
                "status": "Missing"
            })

    for skill in partial:
        key = skill.title()
        if key in UPSKILL_MAP:
            rec = UPSKILL_MAP[key].copy()
            rec["skill"] = skill
            rec["status"] = "Partial"
            rec["priority"] = "Medium"
            recommendations.append(rec)
        else:
            recommendations.append({
                "skill": skill,
                "title": f"Improve {skill}",
                "description": f"Advance your proficiency in {skill}",
                "priority": "Medium",
                "status": "Partial"
            })

    return recommendations

def upskill_cards(recommendations):
    if not recommendations:
        return """
        <div style="
            background:#e8f5e9;
            padding:20px;
            border-radius:14px;
            text-align:center;
            font-size:16px;
            font-weight:600;
            color:#2e7d32;
        ">
            🎉 Excellent! Your skills already match the job requirements.
        </div>
        """, 120

    html = "<div style='display:flex; flex-wrap:wrap; gap:20px;'>"

    for rec in recommendations:
        if rec["priority"] == "High":
            color = "#F44336"
            icon = "🔥"
        else:
            color = "#FFC107"
            icon = "⚡"

        html += f"""
        <div style="
            flex:1 1 320px;
            background:#ffffff;
            border-radius:16px;
            box-shadow:0 10px 24px rgba(0,0,0,0.08);
            padding:18px;
            position:relative;
        ">

            <div style="
                position:absolute;
                left:0;
                top:0;
                bottom:0;
                width:6px;
                background:{color};
                border-radius:16px 0 0 16px;
            "></div>

            <div style="padding-left:10px;">
                <h4 style="margin:0;font-size:18px;">
                    {icon} {rec["skill"]}
                </h4>

                <p style="
                    margin:8px 0;
                    color:#555;
                    font-size:14px;
                ">
                    {rec["description"]}
                </p>

                <span style="
                    background:{color};
                    color:white;
                    padding:4px 12px;
                    border-radius:14px;
                    font-size:12px;
                    font-weight:600;
                ">
                    {rec["priority"]} Priority • {rec["status"]}
                </span>
            </div>
        </div>
        """

    html += "</div>"

    rows = max(1, (len(recommendations) + 2) // 3)  # 3 cards per row
    height = rows * 200 + 40


    return html, height

import base64
def preview_uploaded_document(file):
    if file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            if page.extract_text():
                text += page.extract_text() + "\n\n"

        st.text_area(
            "PDF Content Preview (Parsed)",
            text,
            height=450
        )

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

        st.text_area(
            "DOCX Content Preview",
            text,
            height=450
        )

    else:  # txt
        st.text_area(
            "TXT Content Preview",
            file.read().decode("utf-8"),
            height=450
        )
def extract_text_preserve(file):
    text = ""
    if file.name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for p in doc.paragraphs:
            if p.text.strip():
                text += p.text + "\n"
        text += "\n"
    else:  # txt
        file.seek(0)  # Reset pointer to start
        text = file.read()
        if isinstance(text, bytes):  # decode if needed
            text = text.decode("utf-8", errors="ignore")
        # Normalize line breaks
        text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text


import html
import streamlit as st

def pretty_document_preview(file, max_chars=5000):
    text = extract_text_preserve(file)
    
    # Limit preview to avoid huge scroll
    preview_text = text[:max_chars]
    if len(text) > max_chars:
        preview_text += "\n\n... (truncated for preview)"
    
    # Escape HTML to preserve punctuation
    preview_text = html.escape(preview_text)
    
    # Preserve multiple spaces for indentation
    preview_text = preview_text.replace("  ", "&nbsp;&nbsp;")
    
    # Render nicely
    st.markdown(
        f"""
        <div style="
            background-color:#f7f7f7;
            padding:15px;
            border-radius:8px;
            height:450px;
            overflow:auto;
            white-space: pre-wrap;
            font-family: 'Courier New', monospace;
            font-size:14px;
            line-height:1.5;
            color:#111;
        ">
        {preview_text}
        </div>
        """,
        unsafe_allow_html=True
    )



# --------------------------------------------------
# UPLOAD SECTION
# --------------------------------------------------
st.subheader("📄 Upload Documents")

c1, c2 = st.columns(2)
with c1:
    resume_file = st.file_uploader("Upload Resume", ["pdf", "docx", "txt"])
with c2:
    jd_file = st.file_uploader("Upload Job Description", ["pdf", "docx", "txt"])

if not resume_file or not jd_file:
    st.stop()

# --------------------------------------------------
# PROCESSING
# --------------------------------------------------
resume_raw = extract_text(resume_file)
jd_raw = extract_text(jd_file)

resume_text = clean_text(resume_raw)
jd_text = clean_text(jd_raw)


resume_skills = extract_skills(resume_text)
jd_skills = extract_skills(jd_text)
skill_gap_results = bert_skill_match(resume_skills, jd_skills)

matched = [r["Skill"] for r in skill_gap_results if r["Status"] == "Matched"]
partial = [r["Skill"] for r in skill_gap_results if r["Status"] == "Partial"]
missing = [r["Skill"] for r in skill_gap_results if r["Status"] == "Missing"]


overall_score = round((len(matched) / len(jd_skills)) * 100) if jd_skills else 0


rtech, rsoft = classify(resume_skills)
jtech, jsoft = classify(jd_skills)



# --------------------------------------------------
# MILESTONE 1 – PARSED DOCUMENTS
# --------------------------------------------------
st.divider()
st.header("📌 Milestone 1: Uploaded Document Preview")

c1, c2 = st.columns(2)

with c1:
    st.subheader("📄 Resume Preview")
    if resume_file:
        pretty_document_preview(resume_file)

with c2:
    st.subheader("📄 Job Description Preview")
    if jd_file:
        pretty_document_preview(jd_file)



# --------------------------------------------------
# MILESTONE 2 – SKILL EXTRACTION & PROPORTION
# --------------------------------------------------
st.divider()
st.header("📌 Milestone 2: Skill Extraction & Classification")



c1, c2 = st.columns(2)

with c1:
    st.subheader("Resume Skills")
    
    st.markdown("### 🔧 Technical Skills")
    html = skill_badges(rtech, "#4CAF50")
    components.html(html, height=100, scrolling=True)

    st.markdown("### 🤝 Soft Skills")
    html = skill_badges(rsoft, "#2196F3")
    components.html(html, height=100, scrolling=True)





    fig = px.pie(
        values=[len(rtech), len(rsoft)],
        names=["Technical", "Soft"],
        hole=0.45,
        title="Resume Skill Distribution"
    )
    st.plotly_chart(fig, use_container_width=True)

with c2:
    st.subheader("Job Description Skills")
    st.markdown("### 🔧 Technical Skills")
    html = skill_badges(jtech, "#FF9800")
    components.html(html, height=100, scrolling=True)

    st.markdown("### 🤝 Soft Skills")
    html = skill_badges(jsoft, "#9C27B0")
    components.html(html, height=100, scrolling=True)


    fig = px.pie(
        values=[len(jtech), len(jsoft)],
        names=["Technical", "Soft"],
        hole=0.45,
        title="JD Skill Distribution"
    )
    st.plotly_chart(fig, use_container_width=True)



def skill_summary_card(title, value, color):
    return f"""
    <div style="
        background:#f8f9fc;
        border-radius:14px;
        padding:16px;
        text-align:center;
        box-shadow:0 4px 12px rgba(0,0,0,0.08);
    ">
        <div style="font-size:26px; font-weight:700; color:{color};">
            {value}
        </div>
        <div style="font-size:13px; color:#555; margin-top:4px;">
            {title}
        </div>
    </div>
    """




st.markdown("## 📊 Skill Summaries")

c1, c2 = st.columns(2)

# ---------- Resume ----------
with c1:
    st.markdown("### 📄 Resume Skills")

    resume_summary = f"""
    <div style="
        display:grid;
        grid-template-columns:1fr 1fr;
        gap:12px;
    ">
        {skill_summary_card('Technical Skills', len(rtech), '#4CAF50')}
        {skill_summary_card('Soft Skills', len(rsoft), '#2196F3')}
        {skill_summary_card('Total Skills', len(rtech) + len(rsoft), '#FF9800')}
        {skill_summary_card(
            'Avg Confidence (%)',
            int((len(rtech)/(len(rtech)+len(rsoft))*100)
                if (len(rtech)+len(rsoft)) else 0),
            '#9C27B0'
        )}
    </div>
    """

    components.html(resume_summary, height=260)


# ---------- Job Description ----------
with c2:
    st.markdown("### 📄 Job Description Skills")

    jd_summary = f"""
    <div style="
        display:grid;
        grid-template-columns:1fr 1fr;
        gap:12px;
    ">
        {skill_summary_card('Technical Skills', len(jtech), '#4CAF50')}
        {skill_summary_card('Soft Skills', len(jsoft), '#2196F3')}
        {skill_summary_card('Total Skills', len(jtech) + len(jsoft), '#FF9800')}
        {skill_summary_card(
            'Avg Confidence (%)',
            int((len(jtech)/(len(jtech)+len(jsoft))*100)
                if (len(jtech)+len(jsoft)) else 0),
            '#9C27B0'
        )}
    </div>
    """

    components.html(jd_summary, height=260)


# --------------------------------------------------
# MILESTONE 3 – SKILL GAP INTERFACE (IMAGE 1)
# --------------------------------------------------
st.divider()
st.header("📌 Milestone 3: Skill Embedding Generation And Similarity Computation")

# KPI Cards
k1, k2, k3, k4 = st.columns(4)
k1.metric("Overall Match", f"{overall_score}%")
k2.metric("Matched Skills", len(matched))
k3.metric("Partial Matches", len(partial))
k4.metric("Missing Skills", len(missing))

st.subheader("🧩 Skill Match Breakdown")

c1, c2, c3 = st.columns(3)

with c1:
    st.markdown("### ✅ Matched Skills")
    st.info(comma_list(matched))

with c2:
    st.markdown("### ⚠️ Partial Matched Skills")
    st.warning(comma_list(partial))

with c3:
    st.markdown("### ❌ Missing Skills")
    st.error(comma_list(missing))


st.subheader("Skill Match Chart")
col1, col2 = st.columns([2, 2]) 

rows = []
for skill in jd_skills:
    if skill in matched:
        value = 100
        status = "Matched"
    elif skill in partial:
        value = 50
        status = "Partial"
    else:
        value = 0
        status = "Missing"

    rows.append([skill, value, status])

df = pd.DataFrame(rows, columns=["Skill", "Match", "Status"])


df = pd.DataFrame(rows, columns=["Skill", "Match", "Status"])

fig = px.scatter(
    df,
    x="Match",
    y="Skill",
    color="Status",
    size=[12]*len(df),
    title="Skill Match Status",
    color_discrete_map={
        "Matched": "#4CAF50",
        "Missing": "#F44336",
        "Partial": "#FFC107"
    }
)

fig.update_layout(
    title=dict(
        text="Skill Match Status",
        x=0.5,                 # ✅ CENTER title
        xanchor="center",
        font=dict(color="black", size=20)
    ),
    xaxis=dict(
        title="Match %",
        title_font=dict(color="black", size=14),
        tickfont=dict(color="black", size=12)
    ),
    yaxis=dict(
        title="Skill",
        title_font=dict(color="black", size=14),
        tickfont=dict(color="black", size=12)
    ),
    legend=dict(
        font=dict(color="black", size=12),
        orientation="h",        # ✅ Clean look
        y=-0.25,
        x=0.5,
        xanchor="center"
    ),
    height=380,                # ✅ REDUCE SIZE
    margin=dict(l=140, r=40, t=60, b=40),
    paper_bgcolor="white",
    plot_bgcolor="white"
)
with col1:
    st.plotly_chart(fig, use_container_width=True)


# Donut Overview
fig = px.pie(
    values=[len(matched), len(partial), len(missing)],
    names=["Matched", "Partial", "Missing"],
    hole=0.5,
    title="Skill Match Overview"
)
fig.update_layout(
    title=dict(
        text="Skill Match Overview",
        x=0.5,                 # ✅ center title
        xanchor="center",
        font=dict(color="black", size=20)
    ),
    legend=dict(
        font=dict(color="black"),
        orientation="h",       # clean look
        y=-0.25,
        x=0.5,
        xanchor="center"
    ),
    paper_bgcolor="white",
    plot_bgcolor="white",
    font=dict(color="black"),
    height=300,               # ✅ reduce donut size
    width=300,                # ✅ reduce donut size
    margin=dict(l=20, r=20, t=60, b=20)
)


with col2:
    st.plotly_chart(fig, use_container_width=True)


skills = list(dict.fromkeys(jd_skills + resume_skills))




# --------------------------------------------------
# MILESTONE 4 – DASHBOARD & EXPORT (IMAGE 2)
# --------------------------------------------------

upskill_recommendations = generate_upskilling_recommendations(missing, partial)
upskill_df = pd.DataFrame(upskill_recommendations)

st.divider()
st.header("📌 Milestone 4: Dashboard & Report Export")
c1, c2 = st.columns([5,3])

    # Skill Bar Comparison
treemap_data = []

for skill in skills:
    if skill in resume_skills and skill in jd_skills:
        status = "Matched"
        value = 3
    elif skill in resume_skills:
        status = "Resume Only"
        value = 2
    elif skill in jd_skills:
        status = "JD Only"
        value = 2
    else:
        continue

    treemap_data.append({
        "Category": status,
        "Skill": skill,
        "Value": value
    })

df_tree = pd.DataFrame(treemap_data)



fig = px.treemap(
    df_tree,
    path=["Category", "Skill"],
    values="Value",
    color="Category",
    color_discrete_map={
        "Matched": "#4CAF50",
        "Resume Only": "#2196F3",
        "JD Only": "#FF9800"
    },
    title="Resume vs Job Description Skill Coverage"
)

fig.update_layout(
    title=dict(
        x=0.5,
        xanchor="center",
        font=dict(size=20, color="black")
    ),
    margin=dict(t=60, l=20, r=20, b=20)
)

st.plotly_chart(fig, use_container_width=True)



# --------------------------------------------------
# RADAR CHART – RESUME vs JOB DESCRIPTION (REAL DATA)
# --------------------------------------------------
st.subheader("🕸️ Resume vs Job Description Skill Radar")

# Use union of skills
radar_skills = list(dict.fromkeys(resume_skills + jd_skills))

# Convert skill presence to scores
resume_scores = [
    100 if skill in resume_skills else 0
    for skill in radar_skills
]

job_scores = [
    100 if skill in jd_skills else 0
    for skill in radar_skills
]

# Close radar loop
radar_skills.append(radar_skills[0])
resume_scores.append(resume_scores[0])
job_scores.append(job_scores[0])

# Create radar chart
fig = go.Figure()

fig.add_trace(go.Scatterpolar(
    r=resume_scores,
    theta=radar_skills,
    fill='toself',
    name='Resume',
    line=dict(width=2)
))

fig.add_trace(go.Scatterpolar(
    r=job_scores,
    theta=radar_skills,
    fill='toself',
    name='Job Description',
    line=dict(width=2)
))

fig.update_layout(
    polar=dict(
        radialaxis=dict(
            visible=True,
            range=[0, 100],
            tickvals=[0, 50, 100],
            ticktext=["0", "50", "100"]
        )
    ),
    showlegend=True,
    title="Resume vs Job Description Skill Radar",
    height=500
)

st.plotly_chart(fig, use_container_width=True)



# Upskilling Recommendations
st.subheader("🚀 Upskilling Roadmap Based on Skill Gap")

html, h = upskill_cards(upskill_recommendations)
components.html(html, scrolling=True)






# --------------------------------------------------
# EXPORT SECTION
# --------------------------------------------------
st.subheader("📥 Download Final Report")

def generate_pdf():
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    import tempfile

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")

    # --------------------------------------------------
    # TITLE
    # --------------------------------------------------
    story.append(Paragraph("<b>SkillGap AI – Complete Analysis Report</b>", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"<b>Overall Match Score:</b> {overall_score}%", styles["Normal"]))
    story.append(Spacer(1, 20))

    # --------------------------------------------------
    # MILESTONE 1 – PARSED DOCUMENTS
    # --------------------------------------------------
    #story.append(Paragraph("<b>Milestone 1: Parsed Documents</b>", styles["Heading2"]))
    #story.append(Spacer(1, 8))

    #story.append(Paragraph("<b>Parsed Resume:</b>", styles["Heading3"]))
    #story.append(Paragraph(resume_text[:1500] + "...", styles["Normal"]))
    #story.append(Spacer(1, 10))

    #story.append(Paragraph("<b>Parsed Job Description:</b>", styles["Heading3"]))
    #story.append(Paragraph(jd_text[:1500] + "...", styles["Normal"]))
    #story.append(Spacer(1, 20))
    

    # --------------------------------------------------
    # MILESTONE 2 – SKILL EXTRACTION
    # --------------------------------------------------
    

    story.append(Paragraph(f"<b>Resume Technical Skills:</b> {', '.join(rtech) or 'None'}", styles["Normal"]))
    story.append(Paragraph(f"<b>Resume Soft Skills:</b> {', '.join(rsoft) or 'None'}", styles["Normal"]))
    story.append(Spacer(1, 8))

    story.append(Paragraph(f"<b>JD Technical Skills:</b> {', '.join(jtech) or 'None'}", styles["Normal"]))
    story.append(Paragraph(f"<b>JD Soft Skills:</b> {', '.join(jsoft) or 'None'}", styles["Normal"]))
    story.append(Spacer(1, 20))

    # --------------------------------------------------
    # MILESTONE 3 – SKILL GAP ANALYSIS
    # --------------------------------------------------
   

    story.append(Paragraph("<b>Matched Skills:</b>", styles["Heading3"]))
    for s in matched:
        story.append(Paragraph(f"• {s}", styles["Normal"]))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Partially Matched Skills:</b>", styles["Heading3"]))
    for s in partial:
        story.append(Paragraph(f"• {s}", styles["Normal"]))
    story.append(Spacer(1, 6))

    story.append(Paragraph("<b>Missing Skills:</b>", styles["Heading3"]))
    for s in missing:
        story.append(Paragraph(f"• {s}", styles["Normal"]))
    story.append(Spacer(1, 15))

    # Donut chart (Milestone 3)
    
    # -------- DONUT CHART – PDF SAFE & COLORED --------
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")

    labels = ["Matched", "Partial", "Missing"]
    values = [len(matched), len(partial), len(missing)]

    donut_fig = go.Figure(
        data=[
            go.Pie(
                labels=labels,
                values=values,
                hole=0.55,
                marker=dict(
                    colors=["#4CAF50", "#FFC107", "#F44336"]  # GREEN, YELLOW, RED
                ),
                textinfo="percent+label",
                sort=False
            )
        ]
    )

    donut_fig.update_layout(
        title="Skill Match Overview",
        paper_bgcolor="white",
        plot_bgcolor="white",
        font=dict(color="black"),
        showlegend=True
    )

    pio.write_image(
        donut_fig,
        tmp.name,
        format="png",
        width=600,
        height=400,
        scale=2
    )

    story.append(RLImage(tmp.name, width=300, height=220))



    # --------------------------------------------------
    # MILESTONE 4 – DASHBOARD VISUALS
    # --------------------------------------------------
    

    # Bar chart – PDF export (color safe)
    treemap_data = []

    for skill in skills:
        if skill in resume_skills and skill in jd_skills:
            status = "Matched"
            value = 3
        elif skill in resume_skills:
            status = "Resume Only"
            value = 2
        elif skill in jd_skills:
            status = "JD Only"
            value = 2
        else:
            continue

        treemap_data.append({
            "Category": status,
            "Skill": skill,
            "Value": value
        })

    df_tree = pd.DataFrame(treemap_data)
    # -------- TREEMAP – PDF EXPORT (REQUIRED) --------
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")

    tree_fig = px.treemap(
        df_tree,
        path=["Category", "Skill"],
        values="Value",
        color="Category",
        color_discrete_map={
            "Matched": "#4CAF50",
            "Resume Only": "#2196F3",
            "JD Only": "#FF9800"
        }
    )

    tree_fig.update_layout(
        title=dict(
            text="Resume vs Job Description Skill Coverage",
            x=0.5,
            xanchor="center",
            font=dict(size=18, color="black")
        ),
        paper_bgcolor="white",
        plot_bgcolor="white",
        font=dict(color="black"),
        margin=dict(t=60, l=20, r=20, b=20)
    )

    pio.write_image(
        tree_fig,
        tmp.name,
        format="png",
        width=800,
        height=500,
        scale=2
    )

    story.append(Spacer(1, 12))
    story.append(RLImage(tmp.name, width=400, height=260))
    story.append(Spacer(1, 20))





    # Radar chart – PDF export (color safe)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")

    radar_fig = go.Figure()

    radar_fig.add_trace(go.Scatterpolar(
        r=resume_scores,
        theta=radar_skills,
        fill='toself',
        name='Resume',
        line=dict(color="#2196F3"),
        fillcolor="#90CAF9"
    ))

    radar_fig.add_trace(go.Scatterpolar(
        r=job_scores,
        theta=radar_skills,
        fill='toself',
        name='Job Description',
        line=dict(color="#FF9800"),
        fillcolor="#FFCC80"
    ))

    radar_fig.update_layout(
        polar=dict(radialaxis=dict(range=[0, 100])),
        paper_bgcolor="white",
        plot_bgcolor="white",
        font=dict(color="black"),
        showlegend=True
    )

    pio.write_image(
        radar_fig,
        tmp.name,
        format="png",
        width=600,
        height=500,
        scale=2
    )

    story.append(RLImage(tmp.name, width=300, height=300))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()





csv_data = df.to_csv(index=False).encode("utf-8")

c1, c2 = st.columns(2)

with c1:
    st.download_button(
        "📄 Download PDF",
        generate_pdf(),
        file_name=f"SkillGap_Report_{datetime.now().strftime('%Y%m%d')}.pdf"
    )

with c2:
    st.download_button(
        "📊 Download CSV",
        csv_data,
        file_name="SkillGap_Data.csv"
    )